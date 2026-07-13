from __future__ import annotations

import hashlib
import io
import zipfile
from dataclasses import replace
from datetime import datetime, timezone
from pathlib import Path

import pytest
from docx import Document as WordDocument
from pypdf import PdfReader

from buili_api.services.report_rendering import (
    TEMPLATE_VERSION,
    ReportContext,
    ReportEvidence,
    ReportSource,
    render_docx,
    render_pdf,
    report_sections,
    source_index,
    validate_report_context,
)


ROOT = Path(__file__).resolve().parents[3]
DEMO = ROOT / "buili_demo_evidence"


def _context(kind: str = "evidence_package") -> ReportContext:
    drawing = DEMO / "cooper-residence-E1.1-demo.pdf"
    return ReportContext(
        report_id="rpt_test_bui_1042",
        version=1,
        kind=kind,
        title="Garage GFCI receptacle below required elevation",
        status="draft",
        generated_at=datetime(2026, 7, 12, 17, 25, tzinfo=timezone.utc),
        project_name="Cooper Residence Renovation",
        project_code="CR-2026-017",
        project_address="San Jose, CA",
        issue_number="BUI-1042",
        issue_status="ready_for_review",
        issue_type="quality_defect",
        priority="high",
        classification="unapproved_deviation",
        recommended_action="field_correction_punch",
        evidence_sufficiency="sufficient",
        location="Building: Residence / Space: Garage / Wall: East wall",
        observed_condition="GFCI box centerline measures 12 inches AFF.",
        expected_condition="E1.1 Note 3 requires a minimum 18 inches AFF.",
        difference="Installed centerline is 6 inches below the approved requirement.",
        evidence=[
            ReportEvidence(
                id="evd_demo_measurement",
                kind="measurement",
                title="GFCI centerline tape measurement",
                description="Tape shows approximately 12 inches AFF.",
                transcript="",
                captured_at="2026-07-12T10:18:20-07:00",
                location="Garage / East wall / Entry door",
                image_bytes=(DEMO / "box-elevation-measurement.png").read_bytes(),
            )
        ],
        sources=[
            ReportSource(
                index=1,
                revision_id="rev_demo_e11_2",
                document_title="Electrical Power Plan E1.1",
                sheet_number="E1.1",
                revision="2",
                status="approved",
                page=1,
                quote="Garage receptacle centerline shall be minimum 18 inches AFF.",
                relationship_type="requirement",
                sha256=hashlib.sha256(drawing.read_bytes()).hexdigest(),
            )
        ],
    )


def _pdf_text(payload: bytes) -> str:
    return "\n".join(
        page.extract_text() or "" for page in PdfReader(io.BytesIO(payload)).pages
    )


@pytest.mark.parametrize(
    ("kind", "label"),
    [("rfi", "RFI"), ("punch", "PUNCH"), ("change_event", "CHANGE EVENT")],
)
def test_action_report_kinds_render_source_cited_pdf_and_docx(
    kind: str, label: str
) -> None:
    context = _context(kind)
    pdf = render_pdf(context)
    docx = render_docx(context)

    assert pdf.startswith(b"%PDF-")
    assert label in _pdf_text(pdf)
    assert "BUI-1042" in _pdf_text(pdf)
    assert "E1.1" in _pdf_text(pdf)
    assert len(hashlib.sha256(pdf).hexdigest()) == 64

    assert zipfile.is_zipfile(io.BytesIO(docx))
    editable = WordDocument(io.BytesIO(docx))
    paragraphs = "\n".join(paragraph.text for paragraph in editable.paragraphs)
    assert label in paragraphs
    assert "Source index" in paragraphs
    assert len(editable.tables) >= 5
    assert all(
        shape._inline.docPr.get("descr")
        for shape in editable.inline_shapes
    )
    assert len(hashlib.sha256(docx).hexdigest()) == 64


def test_report_bytes_and_source_manifest_are_reproducible() -> None:
    context = _context()
    first_pdf = render_pdf(context)
    second_pdf = render_pdf(replace(context))
    first_docx = render_docx(context)
    second_docx = render_docx(replace(context))

    assert first_pdf == second_pdf
    assert first_docx == second_docx
    assert PdfReader(io.BytesIO(first_pdf)).metadata.creation_date == datetime(
        2026, 7, 12, 17, 25, tzinfo=timezone.utc
    )
    assert TEMPLATE_VERSION == "buili.project-record.v3"
    assert source_index(context) == [
        {
            "index": 1,
            "revision_id": "rev_demo_e11_2",
            "document_title": "Electrical Power Plan E1.1",
            "sheet_number": "E1.1",
            "revision": "2",
            "status": "approved",
            "page": 1,
            "quote": "Garage receptacle centerline shall be minimum 18 inches AFF.",
            "relationship_type": "requirement",
            "sha256": hashlib.sha256(
                (DEMO / "cooper-residence-E1.1-demo.pdf").read_bytes()
            ).hexdigest(),
        }
    ]
    with zipfile.ZipFile(io.BytesIO(first_docx)) as archive:
        assert archive.namelist() == sorted(archive.namelist())
        assert {item.date_time for item in archive.infolist()} == {
            (1980, 1, 1, 0, 0, 0)
        }


def test_invalid_field_image_does_not_corrupt_report_package() -> None:
    context = _context()
    context.evidence[0].image_bytes = b"not-an-image"
    pdf = render_pdf(context)
    docx = render_docx(context)

    assert "GFCI centerline tape measurement" in _pdf_text(pdf)
    editable = WordDocument(io.BytesIO(docx))
    assert len(editable.inline_shapes) == 0


def test_untrusted_record_text_is_escaped_in_pdf_markup() -> None:
    context = _context()
    context.title = "Pipe A & B <review>"
    context.observed_condition = "Measured A < B & C > D."
    context.evidence[0].title = "Photo <1> & measurement"
    context.sources[0].quote = "Use A&B; do not treat <tag> as markup."

    rendered = _pdf_text(render_pdf(context))
    assert "Pipe A & B <review>" in rendered
    assert "Measured A < B & C > D." in rendered
    assert "Photo <1> & measurement" in rendered
    assert "Use A&B; do not treat <tag> as markup." in rendered


def test_operational_templates_validate_required_fields_and_remove_duplicate_sections() -> None:
    context = _context("rfi")
    assert validate_report_context(context) == ["ball_in_court", "due_date", "question"]
    context.ball_in_court = "Design architect"
    context.due_date = "2026-07-16"
    context.question = context.difference
    context.suggested_answer = "Confirm grid-based layout."

    assert validate_report_context(context) == []
    sections = report_sections(context)
    assert [heading for heading, _ in sections] == [
        "Existing condition",
        "Contract-document conflict",
        "Suggested answer - for review",
    ]
