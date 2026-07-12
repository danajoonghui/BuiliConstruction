from __future__ import annotations

import hashlib
import base64
import io
import zipfile

from pypdf import PdfReader
from docx import Document as WordDocument


async def _upload(client, headers, organization_id, project_id, filename, content_type, data):
    initialized = await client.post(
        "/v1/uploads/init",
        headers=headers,
        json={
            "organization_id": organization_id,
            "project_id": project_id,
            "filename": filename,
            "content_type": content_type,
            "size": len(data),
            "sha256": hashlib.sha256(data).hexdigest(),
        },
    )
    assert initialized.status_code == 201, initialized.text
    upload_id = initialized.json()["data"]["upload_id"]
    stored = await client.put(f"/v1/uploads/{upload_id}/content", headers=headers, content=data)
    assert stored.status_code == 204, stored.text
    completed = await client.post(
        f"/v1/uploads/{upload_id}/complete",
        headers=headers,
        json={"sha256": hashlib.sha256(data).hexdigest()},
    )
    assert completed.status_code == 200, completed.text
    return upload_id


async def test_document_evidence_issue_report_and_rag(client, account):
    headers = account["headers"]
    organizations = (await client.get("/v1/organizations", headers=headers)).json()["data"]
    organization_id = next(item["id"] for item in organizations if item["name"] == "Avery Construction")
    created = await client.post(
        "/v1/projects",
        headers=headers,
        json={
            "organization_id": organization_id,
            "name": "Elm Street Renovation",
            "code": "ELM-001",
            "project_type": "renovation",
            "address": "San Jose, CA",
        },
    )
    assert created.status_code == 201, created.text
    project_id = created.json()["data"]["id"]

    source_bytes = b"E1.1 Electrical Note 3: garage receptacle centerline minimum 18 inches above finished floor."
    source_upload = await _upload(
        client, headers, organization_id, project_id, "E1.1.txt", "text/plain", source_bytes
    )
    document = await client.post(
        f"/v1/projects/{project_id}/documents",
        headers=headers,
        json={"title": "Electrical Plan E1.1", "kind": "drawing", "discipline": "electrical"},
    )
    assert document.status_code == 201, document.text
    document_id = document.json()["data"]["id"]
    revision = await client.post(
        f"/v1/documents/{document_id}/revisions",
        headers=headers,
        json={
            "upload_id": source_upload,
            "revision": "1",
            "status": "approved",
            "sheet_number": "E1.1",
            "process": True,
        },
    )
    assert revision.status_code == 201, revision.text
    revision_id = revision.json()["data"]["revision_id"]
    await client._transport.app.state.services.jobs.queue.join()
    document_detail = await client.get(f"/v1/documents/{document_id}", headers=headers)
    assert "minimum 18 inches" in document_detail.json()["data"]["revisions"][0]["extracted_text"]

    photo = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
    )
    photo_upload = await _upload(client, headers, organization_id, project_id, "context.jpg", "image/jpeg", photo)
    photo_response = await client.post(
        f"/v1/projects/{project_id}/evidence",
        headers=headers,
        json={
            "upload_id": photo_upload,
            "kind": "photo",
            "title": "Garage east wall context",
            "description": "GFCI box visible at rough-in.",
            "location_json": {"space": "Garage", "wall": "East"},
        },
    )
    assert photo_response.status_code == 201, photo_response.text
    photo_id = photo_response.json()["data"]["evidence"]["id"]
    measurement = await client.post(
        f"/v1/projects/{project_id}/evidence",
        headers=headers,
        json={
            "kind": "measurement",
            "title": "Tape measurement",
            "description": "Centerline measures 12 inches AFF.",
            "location_json": {"space": "Garage", "wall": "East"},
            "metadata_json": {
                "measurement": {
                    "observed_value": 12,
                    "unit": "inch",
                    "minimum": 18,
                    "tolerance": 0,
                    "reference": "AFF",
                    "source_revision_id": revision_id,
                }
            },
        },
    )
    assert measurement.status_code == 201, measurement.text
    measurement_id = measurement.json()["data"]["evidence"]["id"]

    issue_response = await client.post(
        f"/v1/projects/{project_id}/issues",
        headers=headers,
        json={
            "title": "Garage receptacle elevation",
            "issue_type": "quality_defect",
            "observed_condition": "Receptacle centerline is 12 inches AFF.",
            "expected_condition": "Approved E1.1 requires centerline minimum 18 inches AFF.",
            "difference": "Installed 6 inches below requirement.",
            "location_json": {"space": "Garage", "wall": "East"},
            "evidence_ids": [photo_id, measurement_id],
            "revision_ids": [revision_id],
        },
    )
    assert issue_response.status_code == 201, issue_response.text
    issue_id = issue_response.json()["data"]["id"]
    relinked_evidence = await client.post(
        f"/v1/issues/{issue_id}/evidence",
        headers=headers,
        json={"evidence_id": photo_id, "relationship_type": "documents"},
    )
    assert relinked_evidence.status_code == 201, relinked_evidence.text
    first_source = await client.post(
        f"/v1/issues/{issue_id}/sources",
        headers=headers,
        json={
            "revision_id": revision_id,
            "quote": "E1.1 Note 3 requires a minimum 18-inch centerline.",
            "relationship_type": "requirement",
        },
    )
    assert first_source.status_code == 201, first_source.text
    second_source = await client.post(
        f"/v1/issues/{issue_id}/sources",
        headers=headers,
        json={
            "revision_id": revision_id,
            "quote": "E1.1 Note 3: garage receptacle centerline minimum 18 inches AFF.",
            "relationship_type": "requirement",
        },
    )
    assert second_source.status_code == 201, second_source.text
    assert second_source.json()["data"]["id"] == first_source.json()["data"]["id"]
    detail_after_relink = await client.get(f"/v1/issues/{issue_id}", headers=headers)
    assert len(detail_after_relink.json()["data"]["sources"]) == 1
    analyzed = await client.post(f"/v1/issues/{issue_id}/analyze", headers=headers)
    assert analyzed.status_code == 202, analyzed.text
    await client._transport.app.state.services.jobs.queue.join()
    issue = (await client.get(f"/v1/issues/{issue_id}", headers=headers)).json()["data"]["issue"]
    assert issue["evidence_sufficiency"] == "sufficient"
    assert issue["classification"] == "unapproved_deviation"
    assert issue["recommended_action"] == "field_correction"

    report = await client.post(
        f"/v1/issues/{issue_id}/reports",
        headers=headers,
        json={"kind": "punch", "title": "Raise garage GFCI before close-in"},
    )
    assert report.status_code == 201, report.text
    report_id = report.json()["data"]["id"]
    report_detail = await client.get(f"/v1/reports/{report_id}", headers=headers)
    assert report_detail.status_code == 200
    report_bytes = await client.get(report_detail.json()["data"]["download_url"], headers=headers)
    assert report_bytes.status_code == 200
    assert report_bytes.headers["content-type"].startswith("application/pdf")
    assert report_bytes.content.startswith(b"%PDF-")
    extracted = "\n".join(
        page.extract_text() or "" for page in PdfReader(io.BytesIO(report_bytes.content)).pages
    )
    assert "Raise garage GFCI" in extracted
    artifact_map = {
        item["format"]: item for item in report_detail.json()["data"]["artifacts"]
    }
    assert set(artifact_map) == {"docx", "json", "pdf"}
    assert all(len(item["sha256"]) == 64 for item in artifact_map.values())
    docx_response = await client.get(artifact_map["docx"]["download_url"], headers=headers)
    assert docx_response.status_code == 200
    assert zipfile.is_zipfile(io.BytesIO(docx_response.content))
    editable = WordDocument(io.BytesIO(docx_response.content))
    editable_text = "\n".join(paragraph.text for paragraph in editable.paragraphs)
    assert "Verification finding" in editable_text
    assert "Source index" in editable_text
    assert len(editable.tables) >= 5
    assert len(editable.inline_shapes) == 1

    blocked = await client.post(f"/v1/reports/{report_id}/approve", headers=headers)
    assert blocked.status_code == 409
    assert blocked.json()["error"]["code"] == "REPORT_APPROVAL_BLOCKED"
    issue_approval = await client.post(f"/v1/issues/{issue_id}/approve", headers=headers)
    assert issue_approval.status_code == 200, issue_approval.text
    approval = await client.post(f"/v1/reports/{report_id}/approve", headers=headers)
    assert approval.status_code == 201, approval.text
    assert approval.json()["data"]["status"] == "approved"
    assert approval.json()["data"]["version"] == 2
    assert approval.json()["data"]["approved_at"]

    search = await client.post(
        f"/v1/projects/{project_id}/search",
        headers=headers,
        json={"query": "garage receptacle 18 inches"},
    )
    assert search.status_code == 200, search.text
    assert search.json()["data"]["hits"]
    answer = await client.post(
        f"/v1/projects/{project_id}/ask",
        headers=headers,
        json={"query": "What is the required garage receptacle elevation?"},
    )
    assert answer.status_code == 200, answer.text
    assert answer.json()["data"]["provider"] == "disabled"
    assert answer.json()["data"]["citations"]
    audit = await client.get(f"/v1/projects/{project_id}/audit", headers=headers)
    actions = {item["action"] for item in audit.json()["data"]}
    assert {"ISSUE_EVIDENCE_LINKED", "ISSUE_SOURCE_LINKED"}.issubset(actions)


async def test_cross_tenant_project_is_forbidden(client, account):
    second = await client.post(
        "/v1/auth/signup?transport=body",
        json={
            "email": "outsider@example.com",
            "password": "AnotherPassword123!",
            "display_name": "Outside User",
            "organization_name": "Outside Org",
        },
    )
    headers = {"Authorization": f"Bearer {second.json()['data']['access_token']}"}
    projects = (await client.get("/v1/projects", headers=account["headers"])).json()["data"]
    private_project = next(item for item in projects if item["code"] == "ELM-001")
    response = await client.get(f"/v1/projects/{private_project['id']}", headers=headers)
    assert response.status_code == 403
    runtime = await client.post(
        f"/v1/projects/{private_project['id']}/spatial-runtime/alignment/solve",
        headers=headers,
        json={
            "plan_graph_id": "private-graph",
            "anchor_pairs": [
                {"plan": [0, 0], "field": [0, 0]},
                {"plan": [1, 0], "field": [1, 0]},
            ],
        },
    )
    assert runtime.status_code == 403
