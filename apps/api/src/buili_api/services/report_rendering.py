"""Deterministic BUILI report renderers.

The template is an original BUILI design.  It intentionally does not depend on a
third-party Canva/marketplace template, remote assets, a browser, or JavaScript.
PDF and DOCX are generated from the same normalized context so an exported report
can be reproduced and its source register can be audited.
"""

from __future__ import annotations

import io
import re
import time
import zipfile
from dataclasses import dataclass, field
from datetime import datetime, timezone
from html import escape as html_escape
from typing import Any

from docx import Document as WordDocument
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image as PillowImage
from reportlab.lib import colors  # type: ignore[import-untyped]
from reportlab.lib.pagesizes import letter  # type: ignore[import-untyped]
from reportlab.lib.styles import (  # type: ignore[import-untyped]
    ParagraphStyle,
    getSampleStyleSheet,
)
from reportlab.lib.units import inch  # type: ignore[import-untyped]
from reportlab.lib.utils import TimeStamp  # type: ignore[import-untyped]
from reportlab.platypus import (  # type: ignore[import-untyped]
    Image,
    KeepTogether,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.pdfgen import canvas as pdf_canvas  # type: ignore[import-untyped]


TEMPLATE_VERSION = "buili.project-record.v3"
GREEN = "#50C878"
GREEN_DARK = "#167A47"
INK = "#17211B"
MUTED = "#667069"
LINE = "#DDE5E0"
PALE = "#EFF9F3"
SOFT = "#F6F8F7"


class _InvariantCanvas(pdf_canvas.Canvas):
    """ReportLab canvas with stable metadata and document identifiers."""

    def __init__(self, *args: Any, **kwargs: Any) -> None:
        generated_at = kwargs.pop("generated_at")
        kwargs["invariant"] = 1
        super().__init__(*args, **kwargs)
        epoch = generated_at.astimezone(timezone.utc).timestamp()
        stamp = TimeStamp(invariant=True)
        stamp.t = epoch
        stamp.lt = time.gmtime(epoch)
        stamp.YMDhms = tuple(stamp.lt)[:6]
        stamp.tzname = "UTC"
        stamp.dhh = 0
        stamp.dmm = 0
        self._doc._timeStamp = stamp


@dataclass(slots=True)
class ReportEvidence:
    id: str
    kind: str
    title: str
    description: str
    transcript: str
    captured_at: str
    location: str
    image_bytes: bytes | None = None


@dataclass(slots=True)
class ReportSource:
    index: int
    revision_id: str
    document_title: str
    sheet_number: str
    revision: str
    status: str
    page: int | None
    quote: str
    relationship_type: str
    sha256: str


@dataclass(slots=True)
class ReportContext:
    report_id: str
    version: int
    kind: str
    title: str
    status: str
    generated_at: datetime
    project_name: str
    project_code: str
    project_address: str
    issue_number: str
    issue_status: str
    issue_type: str
    priority: str
    classification: str
    recommended_action: str
    evidence_sufficiency: str
    location: str
    observed_condition: str
    expected_condition: str
    difference: str
    missing_evidence: list[str] = field(default_factory=list)
    evidence: list[ReportEvidence] = field(default_factory=list)
    sources: list[ReportSource] = field(default_factory=list)
    prepared_by: str = ""
    responsible_party: str = ""
    final_approver: str = ""
    ball_in_court: str = ""
    due_date: str = ""
    question: str = ""
    suggested_answer: str = ""
    official_response: str = ""
    cost_impact: str = ""
    schedule_impact: str = ""
    root_cause: str = ""
    required_action: str = ""
    completion_requirement: str = ""
    origin: str = ""
    change_reason: str = ""
    scope: str = ""
    report_date: str = ""
    weather: str = ""
    work_completed: str = ""
    safety_summary: str = ""
    manpower: list[dict[str, Any]] = field(default_factory=list)
    line_items: list[dict[str, Any]] = field(default_factory=list)
    activity_log: list[dict[str, Any]] = field(default_factory=list)


REPORT_TEMPLATE_SCHEMAS: dict[str, dict[str, tuple[str, ...]]] = {
    "evidence_package": {"required": ("observed_condition", "expected_condition", "difference")},
    "issue_detail": {"required": ("observed_condition", "expected_condition", "difference")},
    "punch": {"required": ("responsible_party", "due_date", "required_action", "completion_requirement")},
    "punch_item": {"required": ("responsible_party", "due_date", "required_action", "completion_requirement")},
    "rfi": {"required": ("ball_in_court", "due_date", "question")},
    "change_event": {"required": ("origin", "change_reason", "scope", "cost_impact", "schedule_impact")},
    "daily_report": {"required": ("report_date", "weather", "work_completed", "safety_summary")},
}


def validate_report_context(context: ReportContext) -> list[str]:
    """Return blocking template omissions before a report can be issued."""

    schema = REPORT_TEMPLATE_SCHEMAS.get(context.kind, REPORT_TEMPLATE_SCHEMAS["issue_detail"])
    missing = [name for name in schema["required"] if not _plain(getattr(context, name, ""), fallback="")]
    if not context.sources and context.kind != "daily_report":
        missing.append("sources")
    if context.evidence_sufficiency.lower() == "insufficient" and not context.missing_evidence:
        missing.append("missing_evidence")
    return missing


def report_sections(context: ReportContext) -> list[tuple[str, str]]:
    """Build non-repeating narrative sections for the selected operational form."""

    kind = context.kind
    if kind in {"punch", "punch_item"}:
        candidates = [
            ("Deficiency", context.observed_condition),
            ("Contract requirement", context.expected_condition),
            ("Required correction", context.required_action or context.recommended_action),
            ("Completion proof", context.completion_requirement),
        ]
    elif kind == "rfi":
        candidates = [
            ("Existing condition", context.observed_condition),
            ("Contract-document conflict", context.difference),
            ("Question", context.question),
            ("Suggested answer - for review", context.suggested_answer),
            ("Official response", context.official_response),
        ]
    elif kind == "change_event":
        candidates = [
            ("Original scope", context.expected_condition),
            ("Changed condition", context.observed_condition),
            ("Basis of change", context.difference),
            ("Commercial position", context.scope),
        ]
    elif kind == "daily_report":
        candidates = [
            ("Work completed", context.work_completed),
            ("Weather and site conditions", context.weather),
            ("Safety and inspections", context.safety_summary),
            ("Open field observation", context.observed_condition),
        ]
    else:
        candidates = [
            ("Observed condition", context.observed_condition),
            ("Required / expected condition", context.expected_condition),
            ("Verified difference", context.difference),
            ("Recommended route", context.required_action or context.recommended_action),
        ]
    seen: set[str] = set()
    output: list[tuple[str, str]] = []
    for heading, body in candidates:
        plain = _plain(body, fallback="")
        key = re.sub(r"\W+", "", plain).lower()
        if not key or key in seen:
            continue
        seen.add(key)
        output.append((heading, plain))
    return output


def _plain(value: Any, *, fallback: str = "Not recorded") -> str:
    if value is None:
        return fallback
    if isinstance(value, dict):
        parts = [f"{key.replace('_', ' ').title()}: {_plain(item, fallback='')}" for key, item in value.items() if item not in (None, "", [], {})]
        return " / ".join(parts) or fallback
    if isinstance(value, list):
        return ", ".join(_plain(item, fallback="") for item in value) or fallback
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", str(value))
    text = " ".join(text.split())
    return text or fallback


def _label(value: str) -> str:
    return _plain(value).replace("_", " ").upper()


def _pdf_plain(value: Any, *, fallback: str = "Not recorded") -> str:
    """Escape untrusted record text before passing it to ReportLab markup."""

    return html_escape(_plain(value, fallback=fallback), quote=True)


def _normalized_image_bytes(value: bytes) -> bytes | None:
    try:
        with PillowImage.open(io.BytesIO(value)) as source_image:
            image = source_image.convert("RGB")
            image.thumbnail((1200, 900))
            output = io.BytesIO()
            image.save(output, format="JPEG", quality=84, optimize=True)
            return output.getvalue()
    except (OSError, ValueError):
        return None


def source_index(context: ReportContext) -> list[dict[str, Any]]:
    return [
        {
            "index": item.index,
            "revision_id": item.revision_id,
            "document_title": item.document_title,
            "sheet_number": item.sheet_number,
            "revision": item.revision,
            "status": item.status,
            "page": item.page,
            "quote": item.quote,
            "relationship_type": item.relationship_type,
            "sha256": item.sha256,
        }
        for item in context.sources
    ]


def _pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "brand": ParagraphStyle("brand", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=15, leading=17, textColor=colors.HexColor(INK)),
        "title": ParagraphStyle("title", parent=base["Heading1"], fontName="Helvetica-Bold", fontSize=23, leading=27, textColor=colors.HexColor(INK), spaceAfter=8),
        "eyebrow": ParagraphStyle("eyebrow", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=7.6, leading=10, textColor=colors.HexColor(GREEN_DARK), spaceAfter=4, characterSpacing=0.8),
        "h2": ParagraphStyle("h2", parent=base["Heading2"], fontName="Helvetica-Bold", fontSize=10.5, leading=13, textColor=colors.HexColor(INK), spaceBefore=15, spaceAfter=6),
        "body": ParagraphStyle("body", parent=base["BodyText"], fontName="Helvetica", fontSize=8.8, leading=13.2, textColor=colors.HexColor(INK), spaceAfter=4),
        "small": ParagraphStyle("small", parent=base["BodyText"], fontName="Helvetica", fontSize=7.3, leading=10.5, textColor=colors.HexColor(MUTED)),
        "label": ParagraphStyle("label", parent=base["BodyText"], fontName="Helvetica-Bold", fontSize=6.8, leading=8.5, textColor=colors.HexColor(MUTED)),
        "white": ParagraphStyle("white", parent=base["BodyText"], fontName="Helvetica-Bold", fontSize=8.3, leading=11, textColor=colors.white),
    }


def _pdf_footer(canvas, doc) -> None:
    canvas.saveState()
    width, _ = letter
    canvas.setStrokeColor(colors.HexColor(LINE))
    canvas.setLineWidth(0.5)
    canvas.line(doc.leftMargin, 0.47 * inch, width - doc.rightMargin, 0.47 * inch)
    canvas.setFont("Helvetica", 6.8)
    canvas.setFillColor(colors.HexColor(MUTED))
    canvas.drawString(doc.leftMargin, 0.29 * inch, "BUILI / SOURCE-CITED CONSTRUCTION RECORD")
    canvas.drawRightString(width - doc.rightMargin, 0.29 * inch, f"Page {doc.page}")
    canvas.restoreState()


def _pdf_table(data, widths, *, header: bool = False, accent: bool = False) -> Table:
    table = Table(data, colWidths=widths, repeatRows=1 if header else 0, hAlign="LEFT")
    commands = [
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LINEBELOW", (0, 0), (-1, -1), 0.35, colors.HexColor(LINE)),
    ]
    if header:
        commands.extend(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(SOFT)),
                ("LINEBELOW", (0, 0), (-1, 0), 0.8, colors.HexColor(LINE)),
            ]
        )
    if accent:
        commands.extend(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(PALE)),
                ("LINEBEFORE", (0, 0), (0, -1), 3, colors.HexColor(GREEN)),
            ]
        )
    table.setStyle(TableStyle(commands))
    return table


def _operational_facts(context: ReportContext) -> list[tuple[str, str]]:
    if context.kind in {"punch", "punch_item"}:
        return [
            ("Responsible trade", context.responsible_party),
            ("Punch manager", context.prepared_by),
            ("Final approver", context.final_approver),
            ("Due date", context.due_date),
        ]
    if context.kind == "rfi":
        return [
            ("Ball in court", context.ball_in_court),
            ("Due date", context.due_date),
            ("Cost impact", context.cost_impact or "To be determined"),
            ("Schedule impact", context.schedule_impact or "To be determined"),
        ]
    if context.kind == "change_event":
        return [
            ("Origin", context.origin),
            ("Change reason", context.change_reason),
            ("Scope", context.scope),
            ("Prepared by", context.prepared_by),
        ]
    if context.kind == "daily_report":
        return [
            ("Report date", context.report_date),
            ("Prepared by", context.prepared_by),
            ("Weather", context.weather),
            ("Safety", context.safety_summary),
        ]
    return [
        ("Priority", _label(context.priority)),
        ("Classification", _label(context.classification)),
        ("Evidence", _label(context.evidence_sufficiency)),
        ("Issue status", _label(context.issue_status)),
    ]


def _pdf_operational_record(context: ReportContext, styles: dict[str, ParagraphStyle]) -> list[Any]:
    facts = _operational_facts(context)
    rows = [
        [Paragraph(_pdf_plain(label.upper()), styles["label"]) for label, _ in facts],
        [Paragraph(_pdf_plain(value), styles["body"]) for _, value in facts],
    ]
    flowables: list[Any] = [_pdf_table(rows, [1.6 * inch] * 4, accent=True)]
    for heading, body in report_sections(context):
        flowables.extend(
            [
                Paragraph(_pdf_plain(heading), styles["h2"]),
                _pdf_table([[Paragraph(_pdf_plain(body), styles["body"])]], [6.4 * inch], accent=heading in {"Required correction", "Question", "Commercial position"}),
            ]
        )
    if context.kind == "change_event" and context.line_items:
        flowables.append(Paragraph("Cost and production line items", styles["h2"]))
        item_rows = [[Paragraph(value, styles["label"]) for value in ("COST CODE", "DESCRIPTION", "QTY / UOM", "ROM")]]
        for item in context.line_items:
            item_rows.append(
                [
                    Paragraph(_pdf_plain(item.get("cost_code")), styles["small"]),
                    Paragraph(_pdf_plain(item.get("description")), styles["small"]),
                    Paragraph(_pdf_plain(f"{item.get('quantity', '-')} {item.get('unit', '')}"), styles["small"]),
                    Paragraph(_pdf_plain(item.get("rom")), styles["small"]),
                ]
            )
        flowables.append(_pdf_table(item_rows, [1.05 * inch, 3.2 * inch, 0.95 * inch, 1.2 * inch], header=True))
    if context.kind == "daily_report" and context.manpower:
        flowables.append(Paragraph("Manpower", styles["h2"]))
        manpower_rows = [[Paragraph(value, styles["label"]) for value in ("COMPANY", "TRADE", "WORKERS", "HOURS")]]
        for item in context.manpower:
            manpower_rows.append(
                [
                    Paragraph(_pdf_plain(item.get("company")), styles["small"]),
                    Paragraph(_pdf_plain(item.get("trade")), styles["small"]),
                    Paragraph(_pdf_plain(item.get("workers")), styles["small"]),
                    Paragraph(_pdf_plain(item.get("hours")), styles["small"]),
                ]
            )
        flowables.append(_pdf_table(manpower_rows, [2.2 * inch, 1.9 * inch, 1.1 * inch, 1.2 * inch], header=True))
    if context.activity_log:
        flowables.append(Paragraph("Activity and chronology", styles["h2"]))
        activity_rows = [[Paragraph("TIME", styles["label"]), Paragraph("ACTOR", styles["label"]), Paragraph("EVENT", styles["label"])]]
        for event in context.activity_log:
            activity_rows.append(
                [
                    Paragraph(_pdf_plain(event.get("timestamp")), styles["small"]),
                    Paragraph(_pdf_plain(event.get("actor")), styles["small"]),
                    Paragraph(_pdf_plain(event.get("event")), styles["small"]),
                ]
            )
        flowables.append(_pdf_table(activity_rows, [1.25 * inch, 1.25 * inch, 3.9 * inch], header=True))
    return flowables


def render_pdf(context: ReportContext) -> bytes:
    """Render a Letter-size, source-cited BUILI issue package."""

    styles = _pdf_styles()
    output = io.BytesIO()
    doc = SimpleDocTemplate(
        output,
        pagesize=letter,
        rightMargin=0.68 * inch,
        leftMargin=0.68 * inch,
        topMargin=0.58 * inch,
        bottomMargin=0.66 * inch,
        title=context.title,
        author="BUILI",
        subject=f"{context.kind} issue package {context.issue_number}",
    )
    status_label = "APPROVED RECORD" if context.status == "approved" else "DRAFT / REVIEW REQUIRED"
    header = Table(
        [[Paragraph("BUIL<span color='#50C878'>I</span>", styles["brand"]), Paragraph(status_label, styles["eyebrow"])]],
        colWidths=[5.25 * inch, 1.15 * inch],
    )
    header.setStyle(
        TableStyle(
            [
                ("ALIGN", (1, 0), (1, 0), "RIGHT"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 9),
                ("LINEBELOW", (0, 0), (-1, -1), 2, colors.HexColor(GREEN)),
            ]
        )
    )
    story: list[Any] = [
        header,
        Spacer(1, 0.2 * inch),
        Paragraph(
            _pdf_plain(
                f"{_label(context.kind)} / {context.issue_number} / V{context.version}"
            ),
            styles["eyebrow"],
        ),
        Paragraph(_pdf_plain(context.title), styles["title"]),
        Paragraph(
            " / ".join(
                _pdf_plain(item)
                for item in (
                    context.project_name,
                    context.project_code,
                    context.location,
                )
            ),
            styles["small"],
        ),
        Spacer(1, 0.16 * inch),
    ]
    story.extend(_pdf_operational_record(context, styles))
    story.append(Paragraph("Evidence register", styles["h2"]))
    evidence_rows = [[Paragraph("TYPE", styles["label"]), Paragraph("EVIDENCE", styles["label"]), Paragraph("LOCATION / CAPTURE", styles["label"])]]
    for evidence_item in context.evidence:
        details = _pdf_plain(
            evidence_item.description or evidence_item.transcript,
            fallback="Linked evidence",
        )
        evidence_rows.append(
            [
                Paragraph(_pdf_plain(_label(evidence_item.kind)), styles["small"]),
                Paragraph(
                    f"<b>{_pdf_plain(evidence_item.title)}</b><br/>{details}",
                    styles["small"],
                ),
                Paragraph(
                    f"{_pdf_plain(evidence_item.location)}<br/>"
                    f"{_pdf_plain(evidence_item.captured_at, fallback='Date not recorded')}",
                    styles["small"],
                ),
            ]
        )
    if len(evidence_rows) == 1:
        evidence_rows.append([Paragraph("-", styles["small"]), Paragraph("No evidence linked.", styles["small"]), Paragraph("-", styles["small"])])
    story.append(
        _pdf_table(
            evidence_rows,
            [1.08 * inch, 3.46 * inch, 1.86 * inch],
            header=True,
        )
    )

    images: list[Any] = []
    for evidence_item in context.evidence:
        if not evidence_item.image_bytes:
            continue
        try:
            normalized = _normalized_image_bytes(evidence_item.image_bytes)
            if not normalized:
                continue
            image = Image(io.BytesIO(normalized), width=1.96 * inch, height=1.35 * inch, kind="proportional")
            cell = Table(
                [
                    [image],
                    [Paragraph(_pdf_plain(evidence_item.title), styles["small"])],
                ],
                colWidths=[2.06 * inch],
            )
            cell.setStyle(
                TableStyle(
                    [
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 0),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                        ("TOPPADDING", (0, 0), (-1, -1), 0),
                        ("BOTTOMPADDING", (0, 0), (-1, 0), 4),
                        ("BOTTOMPADDING", (0, 1), (-1, 1), 0),
                    ]
                )
            )
            images.append(cell)
        except Exception:
            continue
    if images:
        image_rows: list[list[Any]] = []
        for index in range(0, len(images), 3):
            batch = images[index : index + 3]
            image_rows.append(batch + ["" for _ in range(3 - len(batch))])
        gallery = Table(image_rows, colWidths=[2.13 * inch] * 3, hAlign="LEFT")
        gallery.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 0), ("RIGHTPADDING", (0, 0), (-1, -1), 8), ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 3)]))
        story.append(KeepTogether([Paragraph("Field image review", styles["h2"]), gallery]))

    story.extend([Spacer(1, 0.24 * inch), Paragraph("Source index", styles["title"])])
    story.append(
        Paragraph(
            "Every requirement statement in this package must be checked against the exact document revision below. Superseded or unapproved sources cannot authorize work.",
            styles["body"],
        )
    )
    source_rows = [[Paragraph("#", styles["label"]), Paragraph("DOCUMENT / REVISION", styles["label"]), Paragraph("PAGE", styles["label"]), Paragraph("CITED REQUIREMENT", styles["label"])]]
    for report_source in context.sources:
        source_rows.append(
            [
                Paragraph(str(report_source.index), styles["small"]),
                Paragraph(
                    f"<b>{_pdf_plain(report_source.sheet_number or report_source.document_title)}</b> / "
                    f"Rev {_pdf_plain(report_source.revision)} / "
                    f"{_pdf_plain(_label(report_source.status))}"
                    f"<br/>SHA-256: {_pdf_plain(report_source.sha256, fallback='not recorded')}",
                    styles["small"],
                ),
                Paragraph(
                    str(report_source.page) if report_source.page else "-",
                    styles["small"],
                ),
                Paragraph(
                    _pdf_plain(
                        report_source.quote,
                        fallback=(
                            "Linked source; exact clause requires reviewer confirmation."
                        ),
                    ),
                    styles["small"],
                ),
            ]
        )
    if len(source_rows) == 1:
        source_rows.append([Paragraph("-", styles["small"]), Paragraph("No source linked", styles["small"]), Paragraph("-", styles["small"]), Paragraph("Official approval is blocked until a source is linked.", styles["small"])])
    story.append(_pdf_table(source_rows, [0.35 * inch, 2.0 * inch, 0.48 * inch, 3.57 * inch], header=True))
    story.extend(
        [
            Paragraph("Review control", styles["h2"]),
            _pdf_table(
                [
                    [Paragraph("REPORT ID", styles["label"]), Paragraph("VERSION", styles["label"]), Paragraph("GENERATED", styles["label"]), Paragraph("CONTROL", styles["label"])],
                    [Paragraph(_pdf_plain(context.report_id), styles["small"]), Paragraph(str(context.version), styles["small"]), Paragraph(context.generated_at.strftime("%Y-%m-%d %H:%M UTC"), styles["small"]), Paragraph(status_label, styles["eyebrow"])],
                ],
                [2.15 * inch, 0.62 * inch, 1.8 * inch, 1.83 * inch],
                header=True,
            ),
            Spacer(1, 0.18 * inch),
            Paragraph(
                "BUILI assembles a traceable review record. It does not make a design decision, authorize a change, determine contractual entitlement, or replace review by the responsible architect, engineer, contractor, or authority having jurisdiction.",
                styles["small"],
            ),
        ]
    )
    def canvasmaker(*args: Any, **kwargs: Any) -> _InvariantCanvas:
        return _InvariantCanvas(*args, generated_at=context.generated_at, **kwargs)

    doc.build(
        story,
        onFirstPage=_pdf_footer,
        onLaterPages=_pdf_footer,
        canvasmaker=canvasmaker,
    )
    return output.getvalue()


def _word_font(run, *, size: float, color: str = INK, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
    run.bold = bold
    run.italic = italic


def _set_cell(cell, text: str, *, size: float = 8.5, color: str = INK, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.12
    _word_font(paragraph.add_run(_plain(text, fallback="-")), size=size, color=color, bold=bold)


def _cell_fill(cell, value: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), value.lstrip("#"))
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_margins(cell, *, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    width_node = table_pr.first_child_found_in("w:tblW")
    if width_node is None:
        width_node = OxmlElement("w:tblW")
        table_pr.append(width_node)
    width_node.set(qn("w:w"), str(sum(widths_dxa)))
    width_node.set(qn("w:type"), "dxa")
    indent_node = table_pr.first_child_found_in("w:tblInd")
    if indent_node is None:
        indent_node = OxmlElement("w:tblInd")
        table_pr.append(indent_node)
    indent_node.set(qn("w:w"), "120")
    indent_node.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa, strict=True):
            tc_width = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_width)
            tc_width.set(qn("w:w"), str(width))
            tc_width.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)


def _word_table(document, headers: list[str], rows: list[list[str]], widths: list[int], *, accent: bool = False):
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    _table_geometry(table, widths)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for index, value in enumerate(headers):
        _set_cell(table.rows[0].cells[index], value, size=7.5, color=MUTED, bold=True)
        _cell_fill(table.rows[0].cells[index], SOFT)
    for row_index, values in enumerate(rows, start=1):
        for column_index, value in enumerate(values):
            _set_cell(table.rows[row_index].cells[column_index], value)
            if accent:
                _cell_fill(table.rows[row_index].cells[column_index], PALE)
    return table


def _word_operational_record(document, context: ReportContext) -> None:
    facts = _operational_facts(context)
    _word_table(
        document,
        [label for label, _ in facts],
        [[value for _, value in facts]],
        [2340, 2340, 2340, 2340],
        accent=True,
    )
    for heading, body in report_sections(context):
        document.add_heading(heading, level=2)
        _word_table(document, [heading], [[body]], [9360], accent=heading in {"Required correction", "Question", "Commercial position"})
    if context.kind == "change_event" and context.line_items:
        document.add_heading("Cost and production line items", level=2)
        _word_table(
            document,
            ["Cost code", "Description", "Qty / UOM", "ROM"],
            [[
                _plain(item.get("cost_code")),
                _plain(item.get("description")),
                _plain(f"{item.get('quantity', '-')} {item.get('unit', '')}"),
                _plain(item.get("rom")),
            ] for item in context.line_items],
            [1500, 4500, 1500, 1860],
        )
    if context.kind == "daily_report" and context.manpower:
        document.add_heading("Manpower", level=2)
        _word_table(
            document,
            ["Company", "Trade", "Workers", "Hours"],
            [[
                _plain(item.get("company")),
                _plain(item.get("trade")),
                _plain(item.get("workers")),
                _plain(item.get("hours")),
            ] for item in context.manpower],
            [3000, 2700, 1500, 2160],
        )
    if context.activity_log:
        document.add_heading("Activity and chronology", level=2)
        _word_table(
            document,
            ["Time", "Actor", "Event"],
            [[
                _plain(item.get("timestamp")),
                _plain(item.get("actor")),
                _plain(item.get("event")),
            ] for item in context.activity_log],
            [1900, 2000, 5460],
        )


def _word_paragraph(document, text: str, *, size: float = 10.5, color: str = INK, bold: bool = False, before: float = 0, after: float = 6, align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.1
    _word_font(paragraph.add_run(_plain(text)), size=size, color=color, bold=bold)
    return paragraph


def _page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])
    _word_font(run, size=8, color=MUTED)


def _canonicalize_ooxml_zip(payload: bytes) -> bytes:
    """Normalize OOXML ZIP ordering and timestamps for reproducible hashes."""

    source = io.BytesIO(payload)
    output = io.BytesIO()
    with zipfile.ZipFile(source, "r") as archive, zipfile.ZipFile(
        output,
        "w",
        compression=zipfile.ZIP_DEFLATED,
        compresslevel=9,
    ) as normalized:
        for name in sorted(archive.namelist()):
            original = archive.getinfo(name)
            info = zipfile.ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            info.create_system = 0
            info.external_attr = original.external_attr
            normalized.writestr(
                info,
                archive.read(name),
                compress_type=zipfile.ZIP_DEFLATED,
                compresslevel=9,
            )
    return output.getvalue()


def render_docx(context: ReportContext) -> bytes:
    """Render the editable companion using the rfi_response business-brief preset."""

    document = WordDocument()
    properties = document.core_properties
    properties.title = context.title
    properties.subject = f"{context.kind} issue package {context.issue_number}"
    properties.author = "BUILI"
    generated_at = context.generated_at.astimezone(timezone.utc).replace(tzinfo=None)
    properties.created = generated_at
    properties.modified = generated_at
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK.lstrip("#"))
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size, before, after in (("Heading 1", 16, 16, 8), ("Heading 2", 13, 12, 6), ("Heading 3", 12, 8, 4)):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(GREEN_DARK.lstrip("#"))
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.paragraph_format.space_after = Pt(0)
    _word_font(header_paragraph.add_run("BUILI  /  CONSTRUCTION VERIFICATION"), size=8, color=MUTED, bold=True)
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.paragraph_format.space_after = Pt(0)
    footer_paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES
    )
    _word_font(
        footer_paragraph.add_run("SOURCE-CITED REVIEW RECORD"),
        size=7.5,
        color=MUTED,
    )
    _word_font(footer_paragraph.add_run("\tPage "), size=8, color=MUTED)
    _page_field(footer_paragraph)

    _word_paragraph(document, _label(context.kind), size=8.5, color=GREEN_DARK, bold=True, after=3)
    _word_paragraph(document, context.title, size=26, bold=True, after=5)
    _word_paragraph(document, f"{context.project_name} / {context.project_code} / {context.issue_number}", size=10, color=MUTED, after=18)
    _word_operational_record(document, context)
    document.add_heading("Evidence register", level=1)
    evidence_rows = [
        [
            _label(item.kind),
            f"{item.title}\n{item.description or item.transcript or 'Linked evidence'}",
            f"{item.location}\n{item.captured_at or 'Date not recorded'}",
        ]
        for item in context.evidence
    ] or [["-", "No evidence linked", "-"]]
    _word_table(
        document,
        ["Type", "Evidence", "Location / capture"],
        evidence_rows,
        [1600, 4800, 2960],
    )
    image_evidence = [item for item in context.evidence if item.image_bytes]
    if image_evidence:
        document.add_heading("Field image review", level=2)
        for item in image_evidence[:6]:
            raw_image = item.image_bytes
            if raw_image is None:
                continue
            normalized = _normalized_image_bytes(raw_image)
            if not normalized:
                continue
            picture = document.add_paragraph()
            picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
            picture.paragraph_format.space_before = Pt(4)
            picture.paragraph_format.space_after = Pt(3)
            shape = picture.add_run().add_picture(
                io.BytesIO(normalized), width=Inches(4.8)
            )
            shape._inline.docPr.set(
                "descr",
                _plain(f"{item.title}. {item.location}")[:255],
            )
            shape._inline.docPr.set("title", _plain(item.title)[:255])
            _word_paragraph(
                document,
                f"{item.title} / {item.location}",
                size=8.5,
                color=MUTED,
                after=10,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )

    document.add_page_break()
    document.add_heading("Source index", level=1)
    _word_paragraph(
        document,
        "Every requirement statement must be checked against the exact document revision below. Superseded or unapproved sources cannot authorize work.",
        size=9.5,
        color=MUTED,
        after=10,
    )
    source_rows = [
        [
            str(item.index),
            f"{item.sheet_number or item.document_title} / Rev {item.revision} / {_label(item.status)}\nSHA-256: {item.sha256 or 'not recorded'}",
            str(item.page or "-"),
            item.quote or "Linked source; exact clause requires reviewer confirmation.",
        ]
        for item in context.sources
    ] or [["-", "No source linked", "-", "Official approval is blocked until a source is linked."]]
    _word_table(document, ["#", "Document / revision", "Page", "Cited requirement"], source_rows, [420, 2700, 620, 5620])

    document.add_heading("Review control", level=1)
    _word_table(
        document,
        ["Report ID", "Version", "Generated", "Status"],
        [[context.report_id, str(context.version), context.generated_at.strftime("%Y-%m-%d %H:%M UTC"), "APPROVED RECORD" if context.status == "approved" else "DRAFT / REVIEW REQUIRED"]],
        [3000, 900, 2700, 2760],
    )
    _word_paragraph(
        document,
        "BUILI assembles a traceable review record. It does not make a design decision, authorize a change, determine contractual entitlement, or replace review by the responsible architect, engineer, contractor, or authority having jurisdiction.",
        size=8.5,
        color=MUTED,
        before=14,
    )
    output = io.BytesIO()
    document.save(output)
    return _canonicalize_ooxml_zip(output.getvalue())
